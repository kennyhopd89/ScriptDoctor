from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_screenplay_docx(scene_list, output_filename):
    """
    Creates a formatted screenplay .docx file from the scene list.
    
    Format Specs:
    - Font: Courier New, 12pt
    - Margins: Left 1.5", Top 1.0", Bottom 1.0", Right 1.0"
    - Spacing: Single line spacing
    """
    document = Document()
    
    # 1. Page Layout (Margins)
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)
        
    # 2. Default Font Style (Courier New 12pt)
    style = document.styles['Normal']
    font = style.font
    font.name = 'Courier New'
    font.size = Pt(12)
    
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.0
    paragraph_format.space_after = Pt(0)
    
    # 3. Content Loop
    for i, scene in enumerate(scene_list):
        # Add spacing before new scenes (except the first one)
        if i > 0:
            document.add_paragraph("")
            document.add_paragraph("")
            
        # Header: BOLD, UPPERCASE
        header_text = scene.get('header', '').upper()
        p = document.add_paragraph()
        run = p.add_run(header_text)
        run.bold = True
        
        # Content
        content_text = scene.get('content', '')
        # Split by lines to handle formatting better if needed, 
        # but adding as a block is usually fine for basic text. 
        # However, to ensure proper spacing, let's split.
        lines = content_text.split('\n')
        for line in lines:
            document.add_paragraph(line)
            
    # 4. Save
    document.save(output_filename)
    return output_filename
